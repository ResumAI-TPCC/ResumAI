"""
Demo script for RA-24: Resume Parsing
Creates a sample resume and demonstrates the parsing functionality
"""

from pathlib import Path

from docx import Document

from app.services.resume_parser import ResumeParserService

# Create a sample DOCX resume
def create_sample_resume():
    doc = Document()
    
    # Personal Info
    doc.add_heading("John Doe", 0)
    doc.add_paragraph("Email: john.doe@example.com | Phone: +1-555-0123")
    doc.add_paragraph("LinkedIn: https://linkedin.com/in/johndoe")
    
    # Summary
    doc.add_heading("Professional Summary", 1)
    doc.add_paragraph(
        "Experienced software engineer with 5+ years of experience in full-stack development. "
        "Specialized in Python, FastAPI, and cloud technologies."
    )
    
    # Skills
    doc.add_heading("Skills", 1)
    doc.add_paragraph("Python, FastAPI, Docker, Kubernetes, AWS, PostgreSQL, React, TypeScript")
    
    # Work Experience
    doc.add_heading("Work Experience", 1)
    doc.add_paragraph("Senior Software Engineer")
    doc.add_paragraph("Tech Corp Inc.")
    doc.add_paragraph("2020 - Present")
    doc.add_paragraph(
        "• Led development of microservices architecture\n"
        "• Improved system performance by 40%\n"
        "• Mentored junior developers"
    )
    
    doc.add_paragraph("\nSoftware Engineer")
    doc.add_paragraph("StartUp XYZ")
    doc.add_paragraph("2018 - 2020")
    doc.add_paragraph(
        "• Built RESTful APIs using FastAPI\n"
        "• Implemented CI/CD pipelines\n"
        "• Collaborated with cross-functional teams"
    )
    
    # Education
    doc.add_heading("Education", 1)
    doc.add_paragraph("Bachelor of Science in Computer Science")
    doc.add_paragraph("University of Technology")
    doc.add_paragraph("2014 - 2018")
    
    # Save
    sample_path = Path("backend/docs/sample_resume.docx")
    sample_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(sample_path))
    
    return sample_path


def demo_parsing():
    print("=" * 80)
    print("RA-24 Resume Parsing Demo")
    print("=" * 80)
    
    # Create sample resume
    print("\n1. Creating sample resume...")
    resume_path = create_sample_resume()
    print(f"   ✓ Created: {resume_path}")
    
    # Parse resume
    print("\n2. Parsing resume...")
    result = ResumeParserService.parse_file(resume_path, resume_path.name)
    
    # Display results
    print("\n3. Parsing Results:")
    print("-" * 80)
    
    print(f"\n📝 Full Name: {result.full_name}")
    
    print("\n📧 Contact Information:")
    print(f"   Email: {result.contact_info.email}")
    print(f"   Phone: {result.contact_info.phone}")
    print(f"   LinkedIn: {result.contact_info.linkedin}")
    
    if result.summary:
        print("\n💼 Summary:")
        print(f"   {result.summary[:100]}...")
    
    print(f"\n🛠️  Skills ({len(result.skills)}):")
    for skill in result.skills[:10]:
        print(f"   • {skill}")
    
    print(f"\n🎓 Education ({len(result.education)}):")
    for edu in result.education:
        print(f"   • {edu.degree} - {edu.institution}")
    
    print(f"\n💼 Work Experience ({len(result.work_experience)}):")
    for exp in result.work_experience:
        print(f"   • {exp.position} at {exp.company} ({exp.duration})")
    
    print("\n" + "=" * 80)
    print("✓ Demo completed successfully!")
    print("=" * 80)
    
    return result


if __name__ == "__main__":
    demo_parsing()
