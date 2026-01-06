"""
Tests for Resume Parsing Endpoint (RA-24)
"""

from pathlib import Path

import pytest
from fastapi.testclient import TestClient

from app.core.config import settings
from app.main import create_app


@pytest.fixture
def client():
    """Create test client"""
    app = create_app()
    return TestClient(app)


@pytest.fixture
def sample_pdf_path(tmp_path):
    """Create a sample PDF file for testing"""
    pdf_content = b"""%PDF-1.4
1 0 obj
<<
/Type /Catalog
/Pages 2 0 R
>>
endobj
2 0 obj
<<
/Type /Pages
/Kids [3 0 R]
/Count 1
>>
endobj
3 0 obj
<<
/Type /Page
/Parent 2 0 R
/Contents 4 0 R
/MediaBox [0 0 612 792]
>>
endobj
4 0 obj
<<
/Length 44
>>
stream
BT
/F1 12 Tf
100 700 Td
(John Doe) Tj
ET
endstream
endobj
xref
0 5
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000214 00000 n
trailer
<<
/Size 5
/Root 1 0 R
>>
startxref
306
%%EOF"""

    pdf_file = tmp_path / "test_resume.pdf"
    pdf_file.write_bytes(pdf_content)
    return pdf_file


def test_parse_resume_missing_file(client):
    """Test parsing with non-existent file"""
    request_data = {
        "file_id": "test-id-123",
        "storage_path": "storage/resumes/nonexistent/test.pdf",
    }

    response = client.post(f"{settings.api_prefix}/resume/parse", json=request_data)

    assert response.status_code == 404
    assert "not found" in response.json()["detail"].lower()


def test_parse_resume_invalid_format(client, tmp_path):
    """Test parsing with unsupported file format"""
    # Create a text file
    txt_file = tmp_path / "test.txt"
    txt_file.write_text("This is a text file")

    request_data = {
        "file_id": "test-id-456",
        "storage_path": str(txt_file),
    }

    response = client.post(f"{settings.api_prefix}/resume/parse", json=request_data)

    assert response.status_code == 400
    assert "unsupported" in response.json()["detail"].lower()


def test_parse_resume_pdf_success(client, sample_pdf_path):
    """Test successful PDF parsing"""
    request_data = {
        "file_id": "test-id-789",
        "storage_path": str(sample_pdf_path),
    }

    response = client.post(f"{settings.api_prefix}/resume/parse", json=request_data)

    assert response.status_code == 200, f"Error: {response.text}"

    data = response.json()
    assert data["file_id"] == "test-id-789"
    assert data["filename"] == "test_resume.pdf"
    assert data["status"] == "success"
    assert "parsed_data" in data

    parsed_data = data["parsed_data"]
    assert "raw_text" in parsed_data
    assert "contact_info" in parsed_data
    assert "skills" in parsed_data
    assert isinstance(parsed_data["skills"], list)


def test_parse_resume_docx_success(client, tmp_path):
    """Test successful DOCX parsing"""
    # Create a minimal DOCX file
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Smith")
    doc.add_paragraph("jane.smith@example.com")
    doc.add_paragraph("+1-555-1234")
    doc.add_paragraph("Skills: Python, FastAPI, Docker")

    docx_file = tmp_path / "test_resume.docx"
    doc.save(str(docx_file))

    request_data = {
        "file_id": "test-id-999",
        "storage_path": str(docx_file),
    }

    response = client.post(f"{settings.api_prefix}/resume/parse", json=request_data)

    assert response.status_code == 200, f"Error: {response.text}"

    data = response.json()
    assert data["file_id"] == "test-id-999"
    assert data["filename"] == "test_resume.docx"
    assert data["status"] == "success"

    parsed_data = data["parsed_data"]
    assert parsed_data["contact_info"]["email"] == "jane.smith@example.com"
    assert parsed_data["raw_text"]  # Should contain text


def test_parse_resume_schema_validation(client):
    """Test request schema validation"""
    # Missing required fields
    invalid_request = {"file_id": "test-only"}

    response = client.post(
        f"{settings.api_prefix}/resume/parse", json=invalid_request
    )

    assert response.status_code == 422  # Validation error
