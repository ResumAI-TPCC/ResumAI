"""
Tests for Resume Upload Endpoint (RA-24)
"""

import io
import logging
import pytest
from unittest.mock import patch
from fastapi.testclient import TestClient
from google.api_core.exceptions import Forbidden, NotFound
from app.main import create_app
from app.core.config import settings
from app.core.error_templates import GCS_BUCKET_NOT_FOUND, GCS_UPLOAD_FAILED


@pytest.fixture
def sample_pdf_content():
    """Create sample PDF content"""
    return b"""%PDF-1.4
1 0 obj
<<
/Type /Catalog
/Pages 2 0 R
>>
endobj
2 0 obj
<<
/Type /Pages
/Kids [3 0 R]
/Count 1
>>
endobj
3 0 obj
<<
/Type /Page
/Parent 2 0 R
/Contents 4 0 R
/MediaBox [0 0 612 792]
>>
endobj
4 0 obj
<<
/Length 44
>>
stream
BT
/F1 12 Tf
100 700 Td
(John Doe) Tj
ET
endstream
endobj
xref
0 5
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000214 00000 n
trailer
<<
/Size 5
/Root 1 0 R
>>
startxref
306
%%EOF"""


@pytest.fixture
def sample_docx_content():
    """Create sample DOCX content"""
    from docx import Document
    
    doc = Document()
    doc.add_paragraph("Jane Smith")
    doc.add_paragraph("jane.smith@example.com")
    doc.add_paragraph("+1-555-1234")
    doc.add_paragraph("Skills: Python, FastAPI, Docker")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


class FakeBucket:
    def __init__(self):
        self.blobs = {}

    def blob(self, name):
        return FakeBlob(name, self)


class FakeBlob:
    def __init__(self, name, bucket):
        self.name = name
        self.bucket = bucket
        self.metadata = {}

    def upload_from_string(self, content, content_type=None):
        self.bucket.blobs[self.name] = content


class FakeClient:
    def __init__(self) -> None:
        self.bucket_name = None
        self.bucket_obj = FakeBucket()

    def bucket(self, name: str) -> FakeBucket:
        self.bucket_name = name
        return self.bucket_obj


def test_resume_upload_success_pdf():
    app = create_app()
    client = TestClient(app)

    fake_client = FakeClient()
    
    # Patch at usage location: where functions are called, not where they're defined
    with patch("app.services.storage.gcs_service.get_gcs_client", return_value=fake_client), \
         patch("app.services.resume_service.validate_pdf_content", return_value=None), \
         patch.object(settings, "GCS_BUCKET_NAME", "test-bucket"), \
         patch.object(settings, "GCS_OBJECT_PREFIX", "resumes"), \
         patch.object(settings, "GCP_PROJECT_ID", "test-project"):
        
        files = {"file": ("test.pdf", b"%PDF-1.4\nfake\n", "application/pdf")}
        r = client.post(f"{settings.API_PREFIX}/resumes/", files=files)

        assert r.status_code == 201, r.text
        res = r.json()
        assert res["status"] == "ok"
        assert "session_id" in res["data"]
        assert "expire_at" in res["data"]
        # Removed extra fields according to Design Doc 4.2.1
        assert "filename" not in res["data"]
        assert "storage_path" not in res["data"]
        assert fake_client.bucket_name == "test-bucket"
        assert fake_client.bucket_obj.blobs


def test_resume_upload_success_docx():
    app = create_app()
    client = TestClient(app)

    fake_client = FakeClient()
    
    # Use unittest.mock.patch to patch functions in their usage context
    with patch("app.services.storage.gcs_service.get_gcs_client", return_value=fake_client), \
         patch.object(settings, "GCS_BUCKET_NAME", "test-bucket"), \
         patch.object(settings, "GCS_OBJECT_PREFIX", "resumes"), \
         patch.object(settings, "GCP_PROJECT_ID", "test-project"):
        
        files = {"file": ("test.docx", b"fake docx content", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        r = client.post(f"{settings.API_PREFIX}/resumes/", files=files)

        assert r.status_code == 201, r.text
        assert r.json()["data"]["session_id"]


def test_upload_resume_unsupported_format(monkeypatch):
    app = create_app()
    client = TestClient(app)
    
    files = {"file": ("test.exe", b"fake exe content", "application/x-msdownload")}
    response = client.post(f"{settings.API_PREFIX}/resumes/", files=files)
    
    assert response.status_code == 400
    assert "unsupported" in response.json()["detail"].lower()


def test_upload_resume_file_too_large(monkeypatch):
    app = create_app()
    client = TestClient(app)
    
    # Create a file larger than 10MB
    large_content = b"x" * (11 * 1024 * 1024)
    files = {"file": ("large.pdf", large_content, "application/pdf")}
    
    response = client.post(f"{settings.API_PREFIX}/resumes/", files=files)
    
    assert response.status_code == 400
    assert "too large" in response.json()["detail"].lower()


class _FailingBlob(FakeBlob):
    """Blob whose upload raises, to exercise the error classification paths."""

    def __init__(self, name, bucket, error):
        super().__init__(name, bucket)
        self._error = error

    def upload_from_string(self, content, content_type=None):
        raise self._error


class FailingBucket(FakeBucket):
    def __init__(self, error):
        super().__init__()
        self._error = error

    def blob(self, name):
        return _FailingBlob(name, self, self._error)


class FailingClient(FakeClient):
    def __init__(self, error):
        super().__init__()
        self.bucket_obj = FailingBucket(error)


def _upload_pdf_with_client(fake_client):
    app = create_app()
    client = TestClient(app)

    with patch("app.services.storage.gcs_service.get_gcs_client", return_value=fake_client), \
         patch("app.services.resume_service.validate_pdf_content", return_value=None), \
         patch.object(settings, "GCS_BUCKET_NAME", "test-bucket"), \
         patch.object(settings, "GCS_OBJECT_PREFIX", "resumes"), \
         patch.object(settings, "GCP_PROJECT_ID", "test-project"):

        files = {"file": ("test.pdf", b"%PDF-1.4\nfake\n", "application/pdf")}
        return client.post(f"{settings.API_PREFIX}/resumes/", files=files)


def test_upload_reports_missing_bucket_as_configuration_error():
    """A missing bucket is detected from the upload itself, not a pre-flight probe."""
    response = _upload_pdf_with_client(FailingClient(NotFound("bucket gone")))

    assert response.status_code == GCS_BUCKET_NOT_FOUND.code
    assert response.json()["detail"] == GCS_BUCKET_NOT_FOUND.detail


def test_upload_reports_other_failures_as_upload_error():
    response = _upload_pdf_with_client(FailingClient(Forbidden("no access")))

    assert response.status_code == GCS_UPLOAD_FAILED.code
    assert response.json()["detail"] == GCS_UPLOAD_FAILED.detail


def test_upload_failure_is_logged_with_traceback(caplog):
    """The root cause must reach the logs; it used to be swallowed silently."""
    with caplog.at_level(logging.ERROR, logger="app.services.storage.gcs_service"):
        _upload_pdf_with_client(FailingClient(Forbidden("no access")))

    records = [r for r in caplog.records if r.name == "app.services.storage.gcs_service"]
    assert records, "upload failure was not logged"
    assert records[0].exc_info is not None
    assert "no access" in caplog.text
